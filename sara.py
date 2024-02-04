import wikipedia
import pyttsx3
import datetime
import speech_recognition as sr
from googletrans import Translator
import cv2
from docx import Document
from docx.shared import Inches


engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
#print(voices[1].id)
engine.setProperty('voice', voices[1].id)


def speak(audio):
    engine.say(audio)
    engine.runAndWait()


def wishMe():
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("Good Morning.")

    elif hour>=12 and hour<18:
        speak("Good Afternoon.")
         
    else:
        speak("Good Evening.")

    speak("I am sara. please tell me how may I help you")

def translate_speech(src, target_language):
    r = sr.Recognizer()
    translator = Translator()
    
    with sr.Microphone() as source:
        speak("Speak something...")
        audio = r.listen(source)
        
    try:
        text = r.recognize_google(audio, language=src)
        speak(text)
        
        translation = translator.translate(text, src=src, dest=target_language)
        translated_text = translation.text
        
        print("your translation is:")
        speak(translated_text)
        
    except sr.UnknownValueError:
        speak("Unable to recognize speech.")
    except sr.RequestError as e:
        speak("Error:", e)

# Example usage
src = "en"  # English can be altered to any languages 
target_language = "kn"  # spanish can be altered to any languages 


def takeCommand():
    #it takes microphone inut from user and gives string op

    r=sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        r.pause_threshold = 0.9 #noise control
        audio =r.listen(source)
    
    try:
        print("Recognizing...")
        query = r.recognize_google(audio, language='en-in') # type: ignore
        print(f"User said:{query}\n")

    except Exception as e:
       # print(e)
        print("say that again, I did not get you")
        return ":"
    return query

def info_get(prompt):
    getinfo = sr.Recognizer()

    with sr.Microphone() as source:
        print(prompt)
        getinfo.adjust_for_ambient_noise(source)
        audio = getinfo.listen(source, timeout=3)

        try:
            text = getinfo.recognize_google(audio) # type: ignore
            return text
        except sr.UnknownValueError:
            print("couldn't understand audio.")
        except sr.RequestError as e:
            print(f"unable to process yoour request...pls try again {e}")


if __name__ =="__main__":
    wishMe()
    print("\n\n")
    print("shall I capture a picture for the medical record purpose? press 'y' for yes and 'n' for no")
    permit = input(speak("shall I capture a picture for the medical record purpose? press 'y' for yes and 'n' for no"))
    if permit == 'y':
            cap = cv2.VideoCapture(0)

            if not cap.isOpened():
              print("Error: Could not open webcam, permission denied or module aint present")
              exit()

            ret, frame = cap.read()

            cv2.imshow('Captured Image', frame)
            cv2.waitKey(2000)
            cv2.destroyAllWindows()  
            # Save the captured image to a file
            cv2.imwrite('D:\\My codes and programs\\IIT hackathon\\database\\captured_image.jpg', frame)
        # Release the VideoCapture and close the OpenCV window
            cap.release()
            cv2.destroyAllWindows()
        
    while True:
        query = takeCommand().lower()
    #logic for task execution
        
        if 'wikipedia' in query:
            speak("looking into Wikipedia...")
            query = query.replace("Wikipedia","")
            results = wikipedia.summary(query,sentences=2) #try wiki
            speak("ead")
            print(results)
            speak(results)

        elif 'information' in query:
            name = info_get(speak("may I  know your name:"))
            dob = info_get(speak("may I  know your date of birth"))
            gender  = info_get(speak("can you please specify your gender? "))
            country = info_get(speak("which country do you belong to?"))
            city = info_get(speak("may I  know your city "))
            doc = info_get(speak("which doctor do you specifiaclly want to meet :"))
            symptoms = info_get(speak("may I  know the symptoms or discomfort you are facing as keypoints "))
            #add any additional information required...

            #store  this data in a text file
            info = Document()
            info.add_heading("Patient Information", level=2)
            info.add_paragraph(f"Name: {name}")
            info.add_paragraph(f"Date of Birth: {dob}")
            info.add_paragraph(f"Gender: {gender}")
            info.add_paragraph(f"City: {city}")
            info.add_paragraph(f"Doctor to meet: {doc}")
            info.add_paragraph(f"Country: {country}")
            info.add_paragraph(f"Symptoms faced by patient:\n{symptoms}")
            info.add_paragraph("                                               ")
            info.add_heading("Patient image ", level=4)
            info.add_picture('D:\\My codes and programs\\IIT hackathon\\database\\captured_image.jpg', width=Inches(2.0))
            

            info.save("D:\\My codes and programs\\IIT hackathon\\database\\patientinfo.docx")
            speak("User information saved to patient.docx")
            
        elif 'translate for me' in query:
            translate_speech(src, target_language)
        
        elif 'go to sleep' in query:
            quit()
